from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

output_path = r"C:\Users\harsh\Downloads\art-frames-sitemap-flow.docx"

doc = Document()

# Title
p = doc.add_paragraph()
run = p.add_run("Art Frames E-Commerce Site - Sitemap & Flow")
run.bold = True
run.font.size = Pt(16)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")

# Summary
summary = doc.add_paragraph()
summary.add_run("Scope: ").bold = True
summary.add_run("All current routes including API.\n")
summary.add_run("Style: ").bold = True
summary.add_run("Flowchart with decision nodes, plus page list and descriptions.")

doc.add_paragraph("")

# Flowchart (monospace text)
flow_title = doc.add_paragraph()
flow_title.add_run("Flowchart").bold = True

flow = (
    "[Start]\n"
    "   |\n"
    "   v\n"
    "[Home /]\n"
    "   |\\\n"
    "   | \\\n"
    "   v  v\n"
    "[Browse Categories]   [Search]\n"
    "   |\n"
    "   v\n"
    "[Product Detail /product/[id]]\n"
    "   |\n"
    "   v\n"
    "<Decision: Join waitlist?>\n"
    "   | Yes\n"
    "   v\n"
    "[Submit email]\n"
    "   |\n"
    "   v\n"
    "[API POST /api/waitlist]\n"
    "   |\n"
    "   v\n"
    "[Waitlist count updates]\n"
    "   |\n"
    "   v\n"
    "[Waitlist Page /waitlist]\n"
    "\n"
    "No -> Return to /product/[id]\n"
)

flow_para = doc.add_paragraph(flow)
flow_run = flow_para.runs[0]
flow_run.font.name = "Courier New"
flow_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
flow_run.font.size = Pt(10)

# Page list

doc.add_paragraph("")
plist_title = doc.add_paragraph()
plist_title.add_run("Pages & Endpoints").bold = True

pages = [
    ("/", "Homepage. Entry point for browsing categories, sales, and featured items."),
    ("/product/[id]", "Product detail page. Shows art info, price, sizes, and waitlist CTA."),
    ("/waitlist", "User-facing waitlist summary page (client state + count)."),
    ("/api/waitlist", "API endpoint. GET returns count; POST adds a waitlist entry."),
]

for path, desc in pages:
    p = doc.add_paragraph(style=None)
    r = p.add_run(f"{path} - ")
    r.bold = True
    p.add_run(desc)

# Footer

doc.add_paragraph("")
footer = doc.add_paragraph("Generated on February 20, 2026")
footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


doc.save(output_path)
print(output_path)
